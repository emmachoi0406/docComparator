import streamlit as st
import pandas as pd
import difflib
import re
from html import escape
from io import BytesIO
from docx import Document

# ---------- Helper Functions ----------

def highlight_differences(a, b):
    a_words = a.split()
    b_words = b.split()
    matcher = difflib.SequenceMatcher(None, a_words, b_words)
    
    a_out = []
    b_out = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            a_out.extend([escape(word) for word in a_words[i1:i2]])
            b_out.extend([escape(word) for word in b_words[j1:j2]])
        elif tag in ("replace", "delete"):
            a_out.extend([f"<u>{escape(word)}</u>" for word in a_words[i1:i2]])
        if tag in ("replace", "insert"):
            b_out.extend([f"<u>{escape(word)}</u>" for word in b_words[j1:j2]])

    return " ".join(a_out), " ".join(b_out)

def render_html_table(results):
    table_html = """
    <style>
    table {
        width: 100%;
        border-collapse: collapse;
        margin-bottom: 2rem;
        font-size: 15px;
    }
    th, td {
        border: 1px solid #ccc;
        padding: 0.75rem;
        vertical-align: top;
        text-align: left;
    }
    th {
        background-color: #f2f2f2;
    }
    tr.same {
        background-color: #d0f0c0;
    }
    tr.modified {
        background-color: #fff3cd;
    }
    tr.added {
        background-color: #b3d9ff;
    }
    tr.deleted {
        background-color: #ffcccc;
    }
    u {
        text-decoration: underline;
        font-weight: bold;
    }
    </style>
    <table>
        <thead>
            <tr>
                <th>구분</th>
                <th>기존 문구</th>
                <th>개정 문구</th>
            </tr>
        </thead>
        <tbody>
    """

    for row in results:
        status_class = row['Status'].lower()
        table_html += (
            f"<tr class='{status_class}'>"
            f"<td><b>{translate_status(row['Status'])}</b></td>"
            f"<td>{row['Original']}</td>"
            f"<td>{row['Revised']}</td>"
            f"</tr>"
        )

    table_html += "</tbody></table>"
    return table_html

def extract_paragraphs(file):
    doc = Document(file)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def classify_diff(old, new, threshold=0.9):
    if old == new:
        return "Same", old, new
    elif not new:
        return "Deleted", old, "<Deleted>"
    elif not old:
        return "Added", "<New>", new
    else:
        ratio = difflib.SequenceMatcher(None, old, new).ratio()
        if ratio >= threshold:
            return "Modified", old, new
        else:
            return "Modified", old, new

def compare_documents(original_paras, revised_paras):
    sm = difflib.SequenceMatcher(None, original_paras, revised_paras)
    result = []

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for i, j in zip(range(i1, i2), range(j1, j2)):
                result.append({
                    "Status": "Same",
                    "Original": original_paras[i],
                    "Revised": revised_paras[j]
                })

        elif tag == "replace":
            len1 = i2 - i1
            len2 = j2 - j1
            min_len = min(len1, len2)
            for k in range(min_len):
                orig_raw = original_paras[i1 + k]
                rev_raw = revised_paras[j1 + k]

                if orig_raw.strip() == rev_raw.strip():
                    result.append({
                        "Status": "Same",
                        "Original": escape(orig_raw),
                        "Revised": escape(rev_raw)
                    })
                else:
                    orig_diff, rev_diff = highlight_differences(orig_raw, rev_raw)
                    result.append({
                        "Status": "Modified",
                        "Original": orig_diff,
                        "Revised": rev_diff
                    })

            for k in range(min_len, len1):
                result.append({
                    "Status": "Deleted",
                    "Original": original_paras[i1 + k],
                    "Revised": "<Deleted>"
                })
            for k in range(min_len, len2):
                result.append({
                    "Status": "Added",
                    "Original": "<New>",
                    "Revised": revised_paras[j1 + k]
                })

        elif tag == "delete":
            for i in range(i1, i2):
                result.append({
                    "Status": "Deleted",
                    "Original": original_paras[i],
                    "Revised": "<Deleted>"
                })

        elif tag == "insert":
            for j in range(j1, j2):
                result.append({
                    "Status": "Added",
                    "Original": "<New>",
                    "Revised": revised_paras[j]
                })

    return result

def strip_tags(text):
    return re.sub(r'<.*?>', '', text)

def create_docx_report(results):
    doc = Document()
    doc.add_heading("변경 대비표 (수정/신설/삭제 항목)", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '구분'
    hdr_cells[1].text = '기존 문구'
    hdr_cells[2].text = '개정 문구'

    for row in results:
        if row['Status'] == "Same":
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = translate_status(row['Status'])
        row_cells[1].text = strip_tags(row['Original'])
        row_cells[2].text = strip_tags(row['Revised'])

    return doc

def translate_status(status):
    return {
        "Same": "동일",
        "Modified": "일부 수정",
        "Added": "신설",
        "Deleted": "삭제"
    }.get(status, status)

# ---------- Streamlit UI ----------

st.set_page_config(page_title="DOCX 변경 대비표", layout="wide")
st.title("📄 변경 대비표 생성기")

st.markdown("""
Word 문서(.docx) 두 개를 비교하여 문단 단위의 변경 대비표를 생성합니다.  
**기존 문서**와 **개정 문서**를 업로드하세요.
""")

col1, col2 = st.columns(2)

with col1:
    file1 = st.file_uploader("📄 기존 문서 업로드 (.docx)", type="docx")

with col2:
    file2 = st.file_uploader("📝 개정 문서 업로드 (.docx)", type="docx")

if file1 and file2:
    with st.spinner("문단을 분석 중입니다..."):
        original_paragraphs = extract_paragraphs(file1)
        revised_paragraphs = extract_paragraphs(file2)

    st.success("✅ 문단 추출 완료!")
    st.write(f"📄 기존 문서 문단 수: {len(original_paragraphs)}")
    st.write(f"📝 개정 문서 문단 수: {len(revised_paragraphs)}")

    comparison_results = compare_documents(original_paragraphs, revised_paragraphs)

    st.subheader("📊 변경 대비표 (수정된 부분은 밑줄로 강조됨)")
    st.components.v1.html(render_html_table(comparison_results), height=800, scrolling=True)

    # ---- DOCX Export ----
    st.markdown("### 📥 변경 문단 Word 파일로 다운로드")

    filtered_results = [row for row in comparison_results if row['Status'] != "Same"]

    if filtered_results:
        docx_file = create_docx_report(filtered_results)
        buffer = BytesIO()
        docx_file.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="DOCX 파일 다운로드",
            data=buffer,
            file_name="변경_대비표.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.info("변경된 문장이 없습니다. Word 파일이 생성되지 않았습니다.")
else:
    st.warning("두 문서를 모두 업로드해주세요.")



